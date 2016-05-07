from docx import Document
from docx.shared import Inches


document = Document()
document.add_heading('Reporting Document', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)

document.add_picture('monty-truth.png', width=Inches(1.25))

document.add_page_break()

document.save('demo.docx')
