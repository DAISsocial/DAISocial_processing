from .report_generator import ReportGenerator
from docx.shared import Inches


class CheckReport(ReportGenerator):

    def __init__(self):
        super().__init__()

    def create(self, graphic, last_common_media: dict, competitor_media):
        self.document.add_heading('Reporting Document', 0)

        p = self.document.add_paragraph('')
        run = p.add_run('Last media classifier')
        run.add_break()
        run = p.add_run('Most common words that appears together ')
        run.bold = True
        run.add_break()
        for word in last_common_media.get('terms_max'):
            run = p.add_run(str(word))
            run.add_break()
        run = p.add_run('Most common words')
        run.bold = True
        run.add_break()
        for word in last_common_media.get('most_common'):
            run = p.add_run(str(word))
            run.add_break()

        self.document.add_page_break()

        self.document.add_heading('All media graphic by time.', level=1)
        self.document.add_picture('plots/plot.png')
        self.document.add_page_break()

        p = self.document.add_paragraph('')
        run = p.add_run('Competitor media')
        run.bold = True
        run.add_break()
        if len(competitor_media) > 1:
            for key, value in competitor_media:
                run = p.add_run('{} - Summary value: {}'.format(key, value))
                run.add_break()
        else:
            p.add_run('Sorry,there was not enough media by your location....')

        self.document.save('docs/demo.docx')
